#!/usr/bin/env python3
"""Extrae preguntas tipo test desde un .docx y genera JSON.

Uso:
  python extraer_docx.py "ruta/al/examen.docx"
  python extraer_docx.py "ruta/al/examen.docx" -o web/datasets
"""
from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parents[1] / "web" / "datasets"

QUESTION_RE = re.compile(r"^(\d{1,3})\.\s+(.+)$")
OPTION_RE = re.compile(r"^([A-D])\)\s*(.+)$")


def extract(docx_path: Path) -> dict:
    doc = Document(str(docx_path))

    # --- answer key from last table that has 'Resp.' column ---
    answer_key: dict[int, str] = {}
    for table in doc.tables:
        headers = [c.text.strip() for c in table.rows[0].cells]
        if "Resp." in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                # table has columns: Pregunta, Resp., Pregunta, Resp., ...
                for k in range(0, len(cells) - 1, 2):
                    num_str = cells[k]
                    resp = cells[k + 1]
                    if num_str.isdigit() and resp:
                        answer_key[int(num_str)] = resp.upper()
            break

    items: list[dict] = []
    current_q: dict | None = None
    pending_question: tuple[int, str] | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        q_match = QUESTION_RE.match(text)
        if q_match:
            if current_q and len(current_q["answers"]) >= 2:
                items.append(current_q)
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()
            current_q = {
                "page": 1,
                "questionNumber": q_num,
                "question": q_text,
                "answers": [],
            }
            continue

        opt_match = OPTION_RE.match(text)
        if opt_match and current_q is not None:
            label = opt_match.group(1).upper()
            opt_text = opt_match.group(2).strip()
            correct_letter = answer_key.get(current_q["questionNumber"])
            selected = label == correct_letter if correct_letter else False
            current_q["answers"].append({"text": opt_text, "selected": selected})
            continue

    if current_q and len(current_q["answers"]) >= 2:
        items.append(current_q)

    return {
        "sourcePdf": docx_path.name,
        "exportedAt": datetime.now(timezone.utc).isoformat(),
        "totalQuestions": len(items),
        "items": items,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract test questions from .docx")
    parser.add_argument("docx", help="Path to the .docx file")
    parser.add_argument("-o", "--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    args = parser.parse_args()

    docx_path = Path(args.docx)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    data = extract(docx_path)
    stem = docx_path.stem
    out_file = output_dir / f"{stem}.json"
    out_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Extraídas {data['totalQuestions']} preguntas → {out_file}")


if __name__ == "__main__":
    main()
